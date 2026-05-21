# -*- coding: utf-8 -*-
"""
EICU QI 보고서 Word 자동 생성 스크립트
실행: python generate_report.py
출력: EICU_QI_보고서_완성.docx
"""
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 색상 정의 ──────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x4E, 0x79)   # 챕터 제목 (#1F4E79)
BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)   # 표 헤더 (#2E75B6)
BLUE_LIGHT = RGBColor(0xBD, 0xD7, 0xEE)   # 표 헤더 배경 (연한 파랑)
GRAY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)   # 교대 행 배경
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
RED_DARK   = RGBColor(0xC0, 0x00, 0x00)

FONT_KR   = '맑은 고딕'
FONT_NUM  = '맑은 고딕'

# ── 헬퍼: 셀 배경색 ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── 헬퍼: 셀 텍스트 색상·굵기 ─────────────────────────────────────────
def style_cell(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name  = FONT_KR
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    # 동아시아 폰트 지정
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_KR)
    rPr.insert(0, rFonts)

# ── 헬퍼: 셀 수직 정렬 ────────────────────────────────────────────────
def set_cell_valign(cell, align='center'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

# ── 헬퍼: 단락 줄간격 설정 ────────────────────────────────────────────
def set_line_spacing(para, spacing=1.5):
    pPr  = para._p.get_or_add_pPr()
    lSpc = OxmlElement('w:spacing')
    lSpc.set(qn('w:line'),     str(int(spacing * 240)))
    lSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lSpc)

# ── 헬퍼: 단락 간격 설정 ─────────────────────────────────────────────
def set_para_spacing(para, before=0, after=100):
    pPr  = para._p.get_or_add_pPr()
    spc  = OxmlElement('w:spacing')
    spc.set(qn('w:before'), str(before))
    spc.set(qn('w:after'),  str(after))
    try:
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            existing.set(qn('w:before'), str(before))
            existing.set(qn('w:after'),  str(after))
        else:
            pPr.append(spc)
    except Exception:
        pPr.append(spc)

# ── 헬퍼: 폰트 동아시아 설정 ─────────────────────────────────────────
def apply_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_KR
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_KR)
    rFonts.set(qn('w:ascii'),    FONT_KR)
    rFonts.set(qn('w:hAnsi'),    FONT_KR)
    # 기존 rFonts 교체
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)

# ── 헬퍼: 문서 기본 폰트 설정 ────────────────────────────────────────
def set_doc_default_font(doc):
    styles = doc.styles['Normal']
    styles.font.name = FONT_KR
    styles.font.size = Pt(10)
    styles.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)

# ── 헬퍼: 텍스트 단락 추가 ────────────────────────────────────────────
def add_para(doc, text='', size=10, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=60,
             indent_left=0, line_spacing=1.5):
    p    = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before, after)
    set_line_spacing(p, line_spacing)
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    if text:
        run = p.add_run(text)
        apply_font(run, size=size, bold=bold, color=color)
    return p

# ── 헬퍼: 챕터 제목 ───────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=200, after=80)
    set_line_spacing(p, 1.2)
    run = p.add_run(text)
    apply_font(run, size=13 if level == 1 else 11,
               bold=True,
               color=BLUE_DARK if level == 1 else BLUE_MID)
    # 아래 테두리 (level 1만)
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E79')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

# ── 헬퍼: 소제목 ──────────────────────────────────────────────────────
def add_subheading(doc, text):
    return add_heading(doc, text, level=2)

# ── 헬퍼: 표 생성 ─────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    """
    headers: ['열1','열2',...] — 첫 행 파란 헤더
    rows: [['값1','값2',...], ...]
    col_widths: [Cm(x), ...] 또는 None
    """
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = w

    # 헤더 행
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '2E75B6')
        set_cell_valign(cell, 'center')
        style_cell(cell, h, bold=True, color=WHITE, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = 'F8FBFF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_valign(cell, 'center')
            style_cell(cell, str(val), size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

    # 표 아래 간격
    add_para(doc, '', after=80)
    return tbl

# ── 헬퍼: 각주 추가 ───────────────────────────────────────────────────
_fn_counter = [0]

def add_footnote(para, footnote_text):
    """run 다음에 각주 번호 추가 + 각주 내용 (Word footnote XML)"""
    _fn_counter[0] += 1
    n = _fn_counter[0]
    doc = para._p.getparent().getparent().getparent()

    # 각주 참조 run
    run = para.add_run(f'[{n}]')
    run.font.size      = Pt(8)
    run.font.bold      = False
    run.font.color.rgb = BLUE_MID
    run.font.superscript = True

    return n  # 번호만 반환 (실제 각주 XML은 복잡, 여기선 endnote style)

# ── 헬퍼: 구분선 ──────────────────────────────────────────────────────
def add_divider(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(p, 60, 60)

# ── 헬퍼: 글머리 기호 ─────────────────────────────────────────────────
def add_bullet(doc, text, indent=0.5, bullet='•'):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 40)
    set_line_spacing(p, 1.4)
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run(f'{bullet}  {text}')
    apply_font(run, size=10)
    return p

# ── 헬퍼: 번호 글머리 ─────────────────────────────────────────────────
def add_numbered(doc, text, num, indent=0.5):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 50)
    set_line_spacing(p, 1.4)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f'{num}.  {text}')
    apply_font(run, size=10)
    return p

# ── 페이지 여백 설정 ──────────────────────────────────────────────────
def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

# ══════════════════════════════════════════════════════════════════════
# 보고서 생성 시작
# ══════════════════════════════════════════════════════════════════════
doc = Document()
set_margins(doc)

# 기본 스타일 재정의
normal = doc.styles['Normal']
normal.font.name = FONT_KR
normal.font.size = Pt(10)

# ────────────────────────────────────────────────────────────────────
# 표지
# ────────────────────────────────────────────────────────────────────
add_para(doc, '', after=400)  # 여백

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, 0, 120)
run = p.add_run('EICU 물품 재고 불일치 및 코스트 입력 누락')
apply_font(run, size=18, bold=True, color=BLUE_DARK)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p2, 0, 80)
run2 = p2.add_run('개선을 위한 확인체계 개발')
apply_font(run2, size=16, bold=True, color=BLUE_DARK)

add_para(doc, '', after=200)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p3, 0, 60)
run3 = p3.add_run('QI (Quality Improvement) 활동 보고서')
apply_font(run3, size=12, bold=False, color=BLUE_MID)

add_divider(doc)
add_para(doc, '', after=200)

info_rows = [
    ('소 속', '중앙대학교병원 응급중환자실 (EICU)'),
    ('병 상 수', '13병상'),
    ('개 설 일', '2023년 5월'),
    ('작 성 일', datetime.date.today().strftime('%Y년 %m월 %d일')),
]
for label, val in info_rows:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 40)
    r1 = p.add_run(f'{label}:  ')
    apply_font(r1, size=11, bold=True, color=BLUE_DARK)
    r2 = p.add_run(val)
    apply_font(r2, size=11, bold=False)

add_para(doc, '', after=400)
doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# 참고 인용 목록 (문서 전역)
# footnote_refs = []  → 본문에 [¹] 형태 삽입, 마지막에 참고문헌 페이지 추가
# ────────────────────────────────────────────────────────────────────
REFS = {
    1: 'De Bie, A. J. R., et al. (2021). Intelligent checklists improve checklist compliance in the intensive care unit: A prospective before-and-after mixed-method study. British Journal of Anaesthesia, 126(2), 404–414.',
    2: 'Chance, E. A., et al. (2024). The effectiveness of checklists and error reporting systems in enhancing patient safety and reducing medical errors in hospital settings: A narrative review. International Journal of Nursing Sciences, 11(3), 387–398.',
    3: 'Al Ashry, H. S., et al. (2016). Effect of compliance with a nurse-led intensive care unit checklist on clinical outcomes in mechanically and nonmechanically ventilated patients. Journal of Intensive Care Medicine, 31(4), 252–257.',
    4: 'Neve, B. V., & Schmidt, C. P. (2022). Point-of-use hospital inventory management with inaccurate usage capture. Health Care Management Science, 25(1), 126–145.',
}

def ref_mark(doc_para, num):
    """인용 번호를 상위첨자로 삽입"""
    run = doc_para.add_run(f'[{num}]')
    run.font.size        = Pt(8)
    run.font.bold        = False
    run.font.color.rgb   = BLUE_MID
    run.font.superscript = True
    apply_font(run, size=8)

# ────────────────────────────────────────────────────────────────────
# ① 문제 정의
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '① 문제 정의')

add_subheading(doc, '1. EICU(응급 중환자실)란?')
add_para(doc,
    'EICU(Emergency Intensive Care Unit, 응급 중환자실)는 생명을 위협하는 응급 상황의 중증 환자를 '
    '24시간 집중적으로 치료하는 특수 병동입니다. 일반 병동과 달리 한 명의 간호사가 1~2명의 중증 환자를 '
    '전담하며, 기계 호흡(인공호흡기), 지속적 신대체요법(CRRT), 고유량 산소치료(HFNC) 등 복잡하고 '
    '연속적인 처치를 동시에 수행합니다.', after=60)

add_para(doc,
    '중앙대학교병원 EICU는 2023년 5월에 개설된 13병상 규모의 신설 병동으로, '
    '중증 응급 환자를 전담하는 구조입니다. 신설 병동이라는 특성상 업무 프로토콜이 '
    '완전히 정착되지 않은 시점에서 운영되고 있습니다.', after=80)

# 박스 스타일 설명
add_subheading(doc, '2. \'코스트 입력\'이란 무엇인가?')
add_para(doc,
    '간호사는 환자에게 처치를 시행한 뒤, 사용한 물품과 처치 내용을 병원 전산(EMR: Electronic Medical Record, '
    '전자의무기록)에 입력해야 합니다. 이 입력 기록을 \'코스트 입력\'이라고 합니다.', after=40)

# 흐름 설명
box_rows = [
    ['단계', '내용', '결과'],
    ['① 처치 시행', '간호사가 환자에게 물품·처치를 제공', '드레싱, 산소 공급, 약물 투여 등'],
    ['② 코스트 입력', 'EMR에 사용 물품·처치를 기록', '기록이 없으면 다음 단계 불가'],
    ['③ 전산 반영', '청구 시스템에 수량·금액 확정', '불일치 시 청구 누락 발생'],
    ['④ 청구 완료', '병원이 환자·보험에 비용 청구', '병원 수익으로 연결'],
]
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_row = tbl.rows[0]
for i, h in enumerate(['단계', '내용', '결과']):
    set_cell_bg(headers_row.cells[i], '2E75B6')
    style_cell(headers_row.cells[i], h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
data_r = [
    ['① 처치 시행', '간호사가 환자에게 물품·처치를 제공', '드레싱, 산소 공급, 약물 투여 등'],
    ['② 코스트 입력', 'EMR에 사용 물품·처치를 기록', '기록이 없으면 다음 단계 불가'],
    ['③ 전산 반영', '청구 시스템에 수량·금액 확정', '불일치 시 청구 누락 발생'],
    ['④ 청구 완료', '병원이 환자·보험에 비용 청구', '병원 수익으로 연결'],
]
for ri, row_data in enumerate(data_r):
    row = tbl.rows[ri+1]
    bg  = 'F8FBFF' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate(row_data):
        set_cell_bg(row.cells[ci], bg)
        style_cell(row.cells[ci], val, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, '', after=80)

add_subheading(doc, '3. 문제가 되는 점')

# 문제 1
p = add_para(doc, '문제 1.  물품 재고 불일치 (불일치율 68%)', size=10, bold=True, before=40, after=40)
add_para(doc,
    '\'불출증\'은 병원 창고에서 병동으로 물품이 출고되었음을 기록한 문서입니다. '
    '창고에서 나간 물품 수량(불출증 출고량)과 간호사가 실제로 처치에 사용했다고 기록한 수량(처방집계)을 '
    '비교하면, 68%의 경우 두 수치가 일치하지 않습니다. '
    '즉, 물품은 나갔으나 처치 기록이 없거나, 처치는 했으나 물품 출고 기록이 없는 상황이 반복됩니다.',
    after=50)

# 문제 2
p2 = add_para(doc, '문제 2.  코스트 입력 누락 (누락률 90%)', size=10, bold=True, before=20, after=40)
add_para(doc,
    '처치가 이루어진 뒤 EMR에 기록되지 않는 경우가 90%에 달합니다. '
    '특히 산소·HFNC·인공호흡기·CRRT처럼 교대 내내 연속적으로 사용되는 처치는 '
    '사용 시간을 교대 마다 정산해야 하지만, 인수인계 직전까지 미루다가 누락되는 사례가 반복됩니다. '
    '이 경우 병원은 실제 제공한 의료 서비스에 대한 비용을 청구하지 못하게 됩니다.',
    after=60)

# 수치 강조 표
add_table(doc,
    headers=['구분', '측정값', '목표값'],
    rows=[
        ['물품 재고 불일치율', '68%', '10% 이하'],
        ['코스트 입력 누락률', '90%', '10% 이하'],
        ['월 반영 지연 추정액', '₩2,992,752', '₩590,787 이하'],
    ],
    col_widths=[Cm(6), Cm(4), Cm(4)]
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# ② 현황 분석
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '② 현황 분석')

add_subheading(doc, '1. 측정 방법')
add_para(doc, '본 QI 활동은 아래 방법으로 현황을 측정하였습니다.', after=40)
add_bullet(doc, '측정 기간: 28일 관찰, 21일 분석')
add_bullet(doc, '불일치율 산출: 불출증 출고량 ÷ 간호처방집계 청구수량 × 100')
add_bullet(doc, '누락률 산출: 처방코드 없는 항목 수 ÷ 전체 처치 수 × 100')
add_bullet(doc, '연속형 처치 정산율: 이브닝 확인 완료 항목 ÷ 연속형 전체 × 100')
add_para(doc, '', after=60)

add_subheading(doc, '2. 선행연구 근거')
add_para(doc,
    '유사한 문제를 다룬 선행연구에서 체계적인 확인 체계의 효과가 검증된 바 있습니다.',
    after=40)

# 선행연구 표
ref_rows = [
    ['체크리스트 도입 효과',
     'ICU 인텔리전트 체크리스트 도입 후 준수율 73.6% → 100% 향상',
     'De Bie et al., 2021 [1]'],
    ['체크리스트·오류 보고 시스템',
     '의료 오류 및 환자 안전 사고를 유의미하게 감소시킴',
     'Chance et al., 2024 [2]'],
    ['간호사 주도 ICU 체크리스트',
     '준수 시 기계 호흡·비기계 호흡 환자 모두 임상 결과 향상',
     'Al Ashry et al., 2016 [3]'],
    ['Point-of-use 재고 관리',
     '사용 지점 기록 정확도 향상 시 재고 불일치율 유의미하게 감소',
     'Neve & Schmidt, 2022 [4]'],
]
add_table(doc,
    headers=['연구 주제', '주요 결과', '출처'],
    rows=ref_rows,
    col_widths=[Cm(4.5), Cm(7), Cm(4)]
)

add_subheading(doc, '3. 현황 흐름도 (Process Flowchart)')
add_para(doc, '아래는 현재 EICU의 처치 → 입력 → 청구 흐름과 문제 발생 지점을 나타낸 플로우차트입니다.', after=40)

# 흐름도를 표로 구성
flow_table = doc.add_table(rows=1, cols=1)
flow_table.style = 'Table Grid'
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
flow_cell = flow_table.rows[0].cells[0]
set_cell_bg(flow_cell, 'F8FBFF')
flow_cell.width = Cm(14)

flow_lines = [
    ('시작: 처치 시행', 'start'),
    ('↓', 'arrow'),
    ('EMR 코스트 입력 시도 (담당 간호사)', 'step'),
    ('↓', 'arrow'),
    ('입력 완료 여부 확인', 'decision'),
    ('', 'empty'),
    ('YES →  전산 반영 대기  →  청구 완료  ✓', 'ok'),
    ('', 'empty'),
    ('NO  →  입력 실패 원인 (복수 가능)', 'no'),
    ('', 'empty'),
    ('  [원인 ①] 인력 부족 — 중증 환자 동시 처치 중 손 못 뗌', 'cause'),
    ('  [원인 ②] 프로그램 조작 어려움 — 연속형 항목 입력 방식 복잡', 'cause'),
    ('  [원인 ③] 입력 기준 불명확 — 어떤 처치를 언제 입력해야 하는지 미교육', 'cause'),
    ('  [원인 ④] 인수인계 시간 부족 — 교대 마감 직전 확인 불가', 'cause'),
    ('  [원인 ⑤] 확인 체계 부재 — 빠진 항목을 알려주는 시스템 없음', 'cause'),
    ('', 'empty'),
    ('↓ 결과', 'arrow'),
    ('누락 발생  →  재고 불일치 확대  →  미청구 손실  ⚠', 'result'),
]
for text, ftype in flow_lines:
    p = flow_cell.add_paragraph()
    set_para_spacing(p, 0, 20)
    run = p.add_run(text)
    if ftype == 'start':
        apply_font(run, size=10, bold=True, color=BLUE_DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ftype == 'arrow':
        apply_font(run, size=11, bold=False, color=BLUE_MID)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ftype == 'decision':
        apply_font(run, size=10, bold=True, color=RGBColor(0xC0, 0x50, 0x00))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ftype == 'ok':
        apply_font(run, size=10, bold=True, color=RGBColor(0x00, 0x70, 0x40))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif ftype == 'no':
        apply_font(run, size=10, bold=True, color=RED_DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif ftype == 'cause':
        apply_font(run, size=9.5, bold=False, color=RGBColor(0x40, 0x40, 0x40))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1)
    elif ftype == 'result':
        apply_font(run, size=10, bold=True, color=RED_DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        apply_font(run, size=8)

add_para(doc, '', after=80)
doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# ③ 근본원인분석
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '③ 근본원인분석')

add_subheading(doc, '1. 5-Why 분석 (직접 관찰 기반)')
add_para(doc,
    '아래 5-Why 분석은 실습 중 직접 관찰한 내용을 토대로 작성하였습니다.',
    after=40)

add_table(doc,
    headers=['단계', '질문 (Why)', '관찰된 근거', '핵심 원인'],
    rows=[
        ['Why 1', '왜 코스트 입력이 안 되는가?',
         '처치 중에는 입력 시간이 없음', '인력 부족'],
        ['Why 2', '왜 처치 중 입력이 어려운가?',
         '중증 환자 동시 처치로 손을 뗄 수 없음', '고강도 업무 환경'],
        ['Why 3', '왜 나중에 소급 입력도 못 하는가?',
         '연속형 항목(산소·HFNC 등)은 사용 시간 계산이 복잡함', '프로그램 조작 어려움'],
        ['Why 4', '왜 사용 시간 계산이 어려운가?',
         '산소·HFNC를 정확히 몇 시간 사용했는지 기억에 의존', '입력 기준 부재'],
        ['Why 5', '왜 기억에 의존해야 하는가?',
         '빠진 항목을 알려주는 알람·체크리스트 시스템이 없음',
         '확인 체계 부재 (근본원인)'],
    ],
    col_widths=[Cm(1.5), Cm(4.5), Cm(5), Cm(4)]
)

add_subheading(doc, '2. 원인 분류 요약')
add_para(doc, '위 분석을 통해 근본원인은 크게 3개 영역으로 구분됩니다.', after=40)

cause_rows = [
    ['인적 요인', '인력 부족으로 인한 처치 중 입력 불가, 교육 부족'],
    ['시스템 요인', '연속형 항목 입력 방식 복잡, 확인 체계 미비'],
    ['구조적 요인', '처치 유형별 입력 기준 미표준화, 인수인계 프로토콜 부재'],
]
add_table(doc,
    headers=['원인 영역', '세부 내용'],
    rows=cause_rows,
    col_widths=[Cm(4), Cm(11)]
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# ④ 질 향상 개선안 및 실천 방안
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '④ 질 향상 개선안 및 실천 방안 제시')

add_para(doc,
    '확인된 근본원인을 해결하기 위해 두 가지 카테고리의 개선안을 수립하였습니다. '
    '단기 실천 방안은 즉시 적용 가능하며, 장기적으로는 디지털 확인 시스템을 구축합니다.',
    after=60)

# ── 카테고리 A
add_subheading(doc, 'A. 재고관리 개선')

add_table(doc,
    headers=['개선안', '내용', '실천 방법', '시기'],
    rows=[
        ['A-1\n처치 유형별\n입력 기준 표준화',
         '4가지 처치 유형별\n입력 시점 명확화',
         '① 일회성 소모품: 처치 직후 즉시\n'
         '② 연속형 항목(산소·HFNC·인공호흡기·CRRT):\n   교대 종료 30분 전 일괄 정산\n'
         '③ 지속 정주 약물: 교대 후 실사용량 확인\n'
         '④ 수혈: 완료 또는 중단 즉시',
         '즉시'],
        ['A-2\n품목별 적정\n재고(MAX) 설정',
         '과다 적재 기준 품목별 설정 →\n초과 시 자동 경고',
         '디지털 대시보드에서 품목별 MAX 수량 입력\n'
         '기준 초과 시 빨간색 경고 표시',
         '단기'],
        ['A-3\n재고 라벨 부착',
         '물품 보관함에 MAX 수량 시각화',
         '대시보드에서 라벨 인쇄\n→ 물품 보관함에 부착',
         '즉시'],
    ],
    col_widths=[Cm(3), Cm(3.5), Cm(7), Cm(1.5)]
)

# ── 카테고리 B
add_subheading(doc, 'B. 기록/문서 개선')

add_table(doc,
    headers=['개선안', '내용', '실천 방법', '시기'],
    rows=[
        ['B-1\n교대 종료\n체크리스트 도입',
         '8개 항목 교대 마감 전\n클릭형 확인 체크리스트',
         '① 일회성 소모품 전산 반영 완료\n'
         '② 산소/HFNC 이브닝 cost 확인\n'
         '③ 인공호흡기 사용 시간 반영 확인\n'
         '④ CRRT 가동 시간 및 세트 반영 확인\n'
         '⑤ 지속정주약물 사용량 반영 확인\n'
         '⑥ 수혈 완료/중단 반영 확인\n'
         '⑦ 드레싱 세트 반영 확인\n'
         '⑧ 미반영 항목 인수인계 확인',
         '즉시'],
        ['B-2\n듀티 마감 알람',
         '교대 30분 전 미확정 시\n자동 팝업 알람',
         'Day 14:30 / Evening 22:30 / Night 06:30\n'
         '미확인 항목 존재 시 자동 경고 모달 표시',
         '즉시'],
        ['B-3\n인수인계\n디지털 메모',
         '처치 비고사항을\n다음 듀티에 자동 전달',
         '24시간 유효 디지털 포스트잇\n→ 다음 근무자 접속 시 즉시 표시',
         '즉시'],
        ['B-4\n이미지 업로드\n자동 분석',
         '불출증·처방집계\n캡처 사진 업로드\n→ 불일치 자동 감지',
         'AI(Claude Vision) OCR로 항목·수량 자동 추출\n'
         '→ 인라인 수정 테이블 → 분석 확정\n'
         '→ KPI·경고 자동 갱신',
         '단기'],
        ['B-5\n직원 교육\n이수 체계화',
         '신규·기존 직원\n정기 교육 프로그램',
         '코스트 입력 기준 교육 (연 2회 이상)\n'
         '교육 이수율 100% 목표로 KPI 관리',
         '단기'],
    ],
    col_widths=[Cm(3), Cm(3.5), Cm(7), Cm(1.5)]
)

# 디지털 시스템 설명 박스
p = doc.add_paragraph()
set_para_spacing(p, 80, 20)
r = p.add_run('▶  디지털 확인체계 시스템 구성 (EICU QI Monitor)')
apply_font(r, size=10.5, bold=True, color=BLUE_DARK)

box_tbl = doc.add_table(rows=1, cols=1)
box_tbl.style = 'Table Grid'
box_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = box_tbl.rows[0].cells[0]
set_cell_bg(bc, 'EBF3FB')
bc.width = Cm(14)

box_items = [
    '파일 업로드 분석: CSV/XLSX 또는 이미지(JPG·PNG) 업로드 → 불일치 항목 자동 추출',
    'OCR 자동 추출: AI(Claude Vision)가 캡처 화면에서 품목명·수량을 읽고 보정 테이블 제공',
    '처치 유형별 지연율: 8개 처치 유형별 반영 지연율 자동 계산 및 시각화',
    '연속형 자동계산: 산소 Flow 기록 기반 이브닝 cost 자동 산출',
    '주간 Top-5: 최근 7일 누락 다빈도 품목 자동 집계',
    '인수인계 메모: 비고란 입력 → 24시간 다음 근무자에게 자동 전달',
]
for item in box_items:
    bp = bc.add_paragraph()
    set_para_spacing(bp, 0, 30)
    br = bp.add_run(f'  •  {item}')
    apply_font(br, size=9.5)

add_para(doc, '', after=80)
doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# ⑤ 과제해결 성과지표
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '⑤ 과제해결 성과지표')

add_para(doc,
    '개선안 적용 효과를 객관적으로 평가하기 위해 아래 6개 지표를 설정하였습니다. '
    '각 지표는 측정 방법이 명확하며, 대시보드 시스템을 통해 자동 집계됩니다.',
    after=60)

add_table(doc,
    headers=['지표명', '측정 방법', '기준값\n(개선 전)', '목표값\n(개선 후)', '측정 주기'],
    rows=[
        ['물품 반영 일치율',
         '출고량 ÷ 청구량 × 100',
         '32%\n(불일치 68%)',
         '90% 이상', '주 1회'],
        ['코스트 입력 누락률',
         '누락 항목 수 ÷ 전체 처치 수 × 100',
         '90%', '10% 이하', '주 1회'],
        ['교대 종료\n체크리스트 수행률',
         '완료 항목 수 ÷ 8개 × 100',
         '0%\n(미도입)',
         '90% 이상', '매 교대'],
        ['직원 교육 이수율',
         '이수 인원 ÷ 전체 인원 × 100',
         '미측정', '100%', '월 1회'],
        ['연속형 항목\n이브닝 처리율',
         '이브닝 확인 완료 ÷ 연속형 전체 × 100',
         '미측정', '90% 이상', '매 교대'],
        ['반영 지연 추정액',
         '미출고량 × 품목 단가 합계 (월 기준)',
         '₩2,992,752/월',
         '₩590,787/월 이하',
         '월 1회'],
    ],
    col_widths=[Cm(3.5), Cm(4), Cm(2.5), Cm(2.5), Cm(2.5)]
)

add_subheading(doc, '측정 일정')
add_table(doc,
    headers=['측정 시점', '시기', '목적'],
    rows=[
        ['사전 측정 (Baseline)', '개선안 적용 전 4주', '기준값 확정'],
        ['중간 평가', '적용 후 4주', '효과 중간 점검 및 수정'],
        ['사후 평가', '적용 후 8~12주', '최종 성과 확인 및 보고'],
    ],
    col_widths=[Cm(4.5), Cm(4), Cm(6.5)]
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# ⑥ 문제해결 시 기대효과
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '⑥ 문제해결 시 기대효과')

add_para(doc,
    '본 QI 활동의 개선안이 충분히 실천될 경우, 환자·병원·간호사 세 측면에서 다음과 같은 효과가 기대됩니다.',
    after=60)

# 정량 효과
add_subheading(doc, '1. 정량적 기대효과')
add_table(doc,
    headers=['항목', '개선 전', '개선 후 (목표)', '변화량'],
    rows=[
        ['물품 재고 불일치율', '68%', '10% 이하', '↓ 58%p 감소'],
        ['코스트 입력 누락률', '90%', '10% 이하', '↓ 80%p 감소'],
        ['체크리스트 수행률', '0% (미도입)', '90% 이상', '↑ +90%p'],
        ['직원 교육 이수율', '미측정', '100%', '완전 이수'],
        ['반영 지연 추정액 (월)', '₩2,992,752', '₩590,787 이하', '₩2,402,000 절감'],
        ['반영 지연 추정액 (연)', '₩35,913,024', '₩7,089,444 이하', '₩28,823,580 절감'],
    ],
    col_widths=[Cm(4.5), Cm(3), Cm(3.5), Cm(4)]
)

# 정성 효과 — 3개 관점
add_subheading(doc, '2. 정성적 기대효과 — 3개 관점')

# 환자 측면
p = add_para(doc, '환자 측면', size=10.5, bold=True, color=BLUE_DARK, before=60, after=30)
add_bullet(doc, '의료 기록 정확성 향상 → 처치 내역 근거 명확화 → 환자 권리 보호')
add_bullet(doc, '물품 재고 정확도 향상 → 필요 물품의 적시 공급으로 처치 연속성 보장')
add_bullet(doc, '교대 인수인계 정확도 향상 → 처치 누락 없이 다음 근무자에게 연결')
add_para(doc, '', after=20)

# 병원 측면
p2 = add_para(doc, '병원 측면', size=10.5, bold=True, color=BLUE_DARK, before=40, after=30)
add_bullet(doc, '수가 청구 누락 방지 → 연간 약 ₩28,800,000 절감 추정')
add_bullet(doc, 'QI 지표 데이터 자동 수집 → 월별 보고 업무 효율화')
add_bullet(doc, '재고 관리 체계화 → 불필요한 과다 발주 및 폐기 감소')
add_bullet(doc, '장기적: EMR-청구 연계 시스템 도입의 근거 자료로 활용 가능')
add_bullet(doc, '신설 병동(EICU)의 업무 프로토콜 조기 표준화에 기여')
add_para(doc, '', after=20)

# 간호사/수간호사 측면
p3 = add_para(doc, '간호사 및 수간호사 측면', size=10.5, bold=True, color=BLUE_DARK, before=40, after=30)
add_bullet(doc, '교대 마감 체크리스트로 빠진 항목을 구조적으로 확인 → 인지 부담 감소')
add_bullet(doc, '자동 알람(30분 전)으로 누락 항목을 미리 확인 → 인수인계 오류 감소')
add_bullet(doc, '인수인계 디지털 메모로 구두 전달 실수 감소 → 연속 처치 안전성 향상')
add_bullet(doc, '수간호사: 재고 이상 항목 즉시 확인 가능 → 재고 조정 의사결정 효율화')
add_bullet(doc, 'AI OCR 자동 추출로 수기 입력 시간 단축 → 실제 처치에 집중 가능한 시간 확보')

add_para(doc, '', after=60)

# 정성 효과 요약 박스
box_tbl2 = doc.add_table(rows=1, cols=3)
box_tbl2.style = 'Table Grid'
box_tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_3 = ['환자 측면', '병원 측면', '간호사/수간호사 측면']
summaries_3 = [
    '처치 연속성 보장\n기록 정확성 향상\n안전한 의료 제공',
    '연간 ₩28.8M 절감\n재고 체계화\nQI 보고 효율화',
    '인지 부담 감소\n인수인계 오류 감소\n처치 집중 시간 확보',
]
bg_colors = ['E8F5E9', 'E3F2FD', 'FFF8E1']
for ci, (hdr, body, bg) in enumerate(zip(headers_3, summaries_3, bg_colors)):
    cell = box_tbl2.rows[0].cells[ci]
    set_cell_bg(cell, bg)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 60, 30)
    r = p.add_run(hdr)
    apply_font(r, size=10, bold=True, color=BLUE_DARK)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, 10, 60)
    r2 = p2.add_run(body)
    apply_font(r2, size=9)

add_para(doc, '', after=80)
doc.add_page_break()

# ────────────────────────────────────────────────────────────────────
# 참고문헌
# ────────────────────────────────────────────────────────────────────
add_heading(doc, '참고문헌')

for num, ref_text in REFS.items():
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 60)
    set_line_spacing(p, 1.4)
    p.paragraph_format.left_indent       = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r_num = p.add_run(f'[{num}]  ')
    apply_font(r_num, size=9.5, bold=True, color=BLUE_MID)
    r_txt = p.add_run(ref_text)
    apply_font(r_txt, size=9.5)

# ════════════════════════════════════════════════════════════════════
output_path = 'EICU_QI_보고서_완성.docx'
doc.save(output_path)
print(f'[OK] 보고서 생성 완료: {output_path}')
